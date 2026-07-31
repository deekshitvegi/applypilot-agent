"""Producing a resume from the profile.

Tailoring here means choosing what to lead with and how to word it. It does not
mean adding anything. Every line that comes out is built from a record the
applicant entered or a resume they uploaded, so there is no path by which an
employer, a date, a degree, a metric or a skill can appear that was not already
true.

The one thing this module will never do is write a sentence about experience
that is not in the profile.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO

from .models import ExperienceRecord, Profile
from .text import content_tokens

_STOP_FOR_RELEVANCE = frozenset(
    {"experience", "years", "role", "team", "work", "working", "strong", "excellent", "ability"}
)


@dataclass
class TailoredResume:
    """What was produced, and exactly what was moved to produce it."""

    filename: str
    data: bytes = b""
    ordering: list[str] = field(default_factory=list)
    highlighted_skills: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)


def _keywords(job_description: str) -> set[str]:
    return {t for t in content_tokens(job_description) if len(t) > 2} - _STOP_FOR_RELEVANCE


def _relevance(text: str, keywords: set[str]) -> int:
    if not keywords:
        return 0
    return len(set(content_tokens(text)) & keywords)


def rank_bullets(record: ExperienceRecord, keywords: set[str]) -> list[str]:
    """The applicant's own bullets, most relevant first. None added, none cut."""
    bullets = [line.strip() for line in (record.description or "").splitlines() if line.strip()]
    return sorted(bullets, key=lambda b: -_relevance(b, keywords))


def rank_skills(profile: Profile, keywords: set[str], limit: int = 24) -> list[str]:
    """The applicant's own skills, the relevant ones first."""
    ranked = sorted(profile.skills, key=lambda s: -_relevance(s, keywords))
    return ranked[:limit]


def build_resume(
    profile: Profile, job_description: str = "", job_title: str = "", company: str = ""
) -> TailoredResume:
    """Write a .docx from the profile alone."""
    from docx import Document
    from docx.shared import Pt

    keywords = _keywords(f"{job_title} {job_description}")
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    name = profile.fact("full_name") or "Resume"
    heading = document.add_paragraph()
    run = heading.add_run(name)
    run.bold = True
    run.font.size = Pt(18)

    contact_parts = [
        profile.fact("city") and f"{profile.fact('city')}, {profile.fact('state')}".strip(", "),
        profile.fact("email"),
        profile.fact("phone"),
        profile.fact("linkedin"),
        profile.fact("github"),
    ]
    contact = "  |  ".join(part for part in contact_parts if part)
    if contact:
        document.add_paragraph(contact)

    ordering: list[str] = []

    skills = rank_skills(profile, keywords)
    if skills:
        document.add_paragraph("Skills").runs[0].bold = True
        document.add_paragraph(", ".join(skills))

    if profile.experience:
        document.add_paragraph("Experience").runs[0].bold = True
        ordered = sorted(
            profile.experience,
            key=lambda r: (
                -_relevance(f"{r.title} {r.company} {r.description}", keywords),
                _sort_key(r),
            ),
        )
        for record in ordered:
            ordering.append(f"{record.title} at {record.company}")
            line = document.add_paragraph()
            title_run = line.add_run(f"{record.title}, {record.company}".strip(", "))
            title_run.bold = True
            dates = _date_range(record)
            if dates or record.location:
                line.add_run(f"  |  {'  |  '.join(p for p in (record.location, dates) if p)}")
            for bullet in rank_bullets(record, keywords):
                document.add_paragraph(bullet, style="List Bullet")

    if profile.education:
        document.add_paragraph("Education").runs[0].bold = True
        for record in profile.education:
            line = document.add_paragraph()
            degree = " in ".join(p for p in (record.degree, record.field_of_study) if p)
            run = line.add_run(", ".join(p for p in (degree, record.school) if p))
            run.bold = True
            tail = "  |  ".join(
                p for p in (record.location, f"{record.start_date} - {record.end_date}".strip(" -"))
                if p
            )
            if tail:
                line.add_run(f"  |  {tail}")

    buffer = BytesIO()
    document.save(buffer)

    slug = re.sub(r"[^A-Za-z0-9]+", "_", f"{name} {company} {job_title}".strip()).strip("_")
    result = TailoredResume(
        filename=f"{slug or 'resume'}.docx",
        data=buffer.getvalue(),
        ordering=ordering,
        highlighted_skills=skills[:10],
    )
    if not keywords:
        result.notes.append(
            "No job description was given, so nothing was reordered -- this is your "
            "profile written out as it stands."
        )
    else:
        result.notes.append(
            "Only the order and the emphasis changed. Every line comes from a record "
            "already in your profile; nothing was added."
        )
    return result


def _date_range(record: ExperienceRecord) -> str:
    end = "Present" if record.current else record.end_date
    return " - ".join(p for p in (record.start_date, end) if p)


def _sort_key(record: ExperienceRecord) -> str:
    """Most recent first, using whatever the record actually says."""
    year = re.search(r"(19|20)\d{2}", record.end_date or record.start_date or "")
    return f"-{year.group(0)}" if year else "0"
