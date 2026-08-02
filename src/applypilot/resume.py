"""Reading a resume into structured records.

Everything here is extraction. Nothing is inferred, rounded, tidied up or filled
in: a field the document does not state comes back empty and onboarding asks for
it. A resume is the applicant's account of their own history and this code has
no business improving on it.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from .matching import expand_state
from .models import EducationRecord, ExperienceRecord

MONTHS = (
    "january|february|march|april|may|june|july|august|september|october|november|december"
    "|jan|feb|mar|apr|jun|jul|aug|sep|sept|oct|nov|dec"
)

_DATE = rf"(?:(?:{MONTHS})\.?\s*)?(?:\d{{1,2}}[/-])?(?:19|20)\d{{2}}"
_RANGE = re.compile(
    rf"(?P<start>{_DATE})\s*(?:[-–—]|to|until)\s*(?P<end>{_DATE}|present|current|now|ongoing)",
    re.IGNORECASE,
)

_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}")
_LINKEDIN = re.compile(r"(?:https?://)?(?:[\w-]+\.)?linkedin\.com/in/[\w%-]+/?", re.IGNORECASE)
_GITHUB = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w-]+/?", re.IGNORECASE)
_WEBSITE = re.compile(r"https?://[\w.-]+\.[a-z]{2,}(?:/\S*)?", re.IGNORECASE)

# Longest spellings first so "M.B.A." is not read as "M.A.", and a trailing dot
# is kept: "M.S." is how it is written on the document.
_DEGREE = re.compile(
    r"\b(ph\.?d\.?|doctorate|m\.?b\.?a\.?|m\.?tech\.?|m\.?eng\.?|m\.?sc?\.?|m\.?a\.?|"
    r"master'?s?|b\.?tech\.?|b\.?eng\.?|b\.?sc?\.?|b\.?a\.?|bachelor'?s?|"
    r"associate'?s?|diploma)(?![a-z])",
    re.IGNORECASE,
)
# Word-boundary on both sides and no hyphen: a thesis about "College-Level
# MCQs" is not an education entry, and it was read as one.
_SCHOOL = re.compile(
    r"(?<![\w-])(university|universitat|college|institute|institut|school|academy|"
    r"polytechnic|iit|nit|iiit)(?![\w-])",
    re.IGNORECASE,
)
_FIELD_OF_STUDY = re.compile(r"\b(?:in|of)\s+(.+)$", re.IGNORECASE)

#: Separators a resume uses between the parts of one heading line.
_CHUNK = re.compile(r"\s*[\t|·•]\s*|\s{3,}")
_LOCATION = re.compile(r"^[A-Z][\w.' -]{1,28},\s*(?:[A-Z]{2}|[A-Z][a-z]+(?: [A-Z][a-z]+)?)$")

SECTION_WORDS = {
    "experience": (
        "experience", "work experience", "professional experience", "employment",
        "employment history", "work history", "career history", "professional background",
        "research experience", "relevant experience", "industry experience",
        "internships", "internship experience", "leadership experience",
    ),
    "education": ("education", "academic background", "academics", "qualifications"),
    "skills": ("skills", "technical skills", "core skills", "technologies", "skills and tools"),
    "projects": ("projects", "selected projects", "personal projects"),
    "summary": ("summary", "profile", "objective", "about", "professional summary"),
    "certifications": ("certifications", "certificates", "licenses", "awards", "publications"),
}


@dataclass
class ResumeExtract:
    """What the document actually said."""

    name: str = ""
    email: str = ""
    phone: str = ""
    linkedin: str = ""
    github: str = ""
    website: str = ""
    location: str = ""
    education: list[EducationRecord] = field(default_factory=list)
    experience: list[ExperienceRecord] = field(default_factory=list)
    skills: list[str] = field(default_factory=list)
    lines: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)

    def as_facts(self) -> dict[str, str]:
        """Only the facts the document stated. Absent stays absent."""
        found = {
            "full_name": self.name,
            "email": self.email,
            "phone": self.phone,
            "linkedin": self.linkedin,
            "github": self.github,
            "website": self.website,
        }
        first, _, last = self.name.partition(" ")
        if last:
            found["first_name"] = first
            found["last_name"] = last.split()[-1]
        if self.location and "," in self.location:
            # "Austin, TX" from a resume, "Austin, Texas, United States" from a
            # LinkedIn export. Splitting on the first comma alone made the state
            # "Texas, United States".
            parts = [part.strip() for part in self.location.split(",") if part.strip()]
            if parts:
                found["city"] = parts[0]
            if len(parts) >= 2:
                found["state"] = expand_state(parts[1])
            if len(parts) >= 3:
                found["country"] = parts[2]
        return {key: value for key, value in found.items() if value}


@dataclass
class Line:
    text: str
    bullet: bool = False


def _is_bullet(paragraph) -> bool:
    """Whether Word considers this paragraph a list item.

    Word strips the bullet glyph out of the text, so a bullet and a heading read
    identically. Telling them apart matters: one line of achievements once ended
    up recorded as the name of an employer.
    """
    style = (getattr(paragraph.style, "name", "") or "").lower()
    if "list" in style or "bullet" in style:
        return True
    try:
        return paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    except AttributeError:
        return False


def read_docx(source: Path | bytes) -> list[Line]:
    """Every line of a .docx, paragraphs and table cells alike."""
    from docx import Document

    handle = BytesIO(source) if isinstance(source, bytes) else str(source)
    document = Document(handle)

    lines: list[Line] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.replace("\xa0", " ").rstrip()
        if text.strip():
            lines.append(Line(text=text, bullet=_is_bullet(paragraph)))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for part in cell.text.replace("\xa0", " ").splitlines():
                    stripped = part.rstrip()
                    if stripped.strip() and stripped not in [ln.text for ln in lines[-3:]]:
                        lines.append(Line(text=stripped))
    return lines


def read_docx_lines(source: Path | bytes) -> list[str]:
    return [line.text.strip() for line in read_docx(source)]


def _section_of(line: str) -> str:
    stripped = re.sub(r"[^a-z ]", " ", line.lower()).strip()
    stripped = re.sub(r"\s+", " ", stripped)
    if len(stripped) > 40:
        return ""
    for name, words in SECTION_WORDS.items():
        if stripped in words:
            return name
    return ""


def _split_sections(lines: list[Line]) -> dict[str, list[Line]]:
    sections: dict[str, list[Line]] = {"header": []}
    current = "header"
    for line in lines:
        name = _section_of(line.text)
        if name:
            current = name
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return sections


def _chunks(text: str) -> list[str]:
    """A heading line broken at the separators a resume actually uses."""
    return [part.strip(" ,|·•\t") for part in _CHUNK.split(text) if part.strip(" ,|·•\t")]


def _looks_like_location(text: str) -> bool:
    return bool(_LOCATION.match(text.strip())) and len(text) <= 40


def _tidy_date(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip(" .,-–—")


def _strip_dates(text: str) -> str:
    return _RANGE.sub("", text).strip(" |,·•-–—\t")


def _parse_experience(lines: list[Line]) -> list[ExperienceRecord]:
    """One record per dated heading, with the bullets under it as description."""
    records: list[ExperienceRecord] = []
    pending: list[str] = []

    for line in lines:
        text = line.text.strip()
        match = _RANGE.search(text)

        if line.bullet or (not match and (len(text) > 150 or text.endswith("."))):
            if records:
                records[-1].description = f"{records[-1].description}\n{text}".strip()
            continue

        if not match:
            pending = _chunks(text)
            continue

        ongoing = bool(re.match(r"present|current|now|ongoing", match.group("end"), re.IGNORECASE))
        record = ExperienceRecord(
            start_date=_tidy_date(match.group("start")),
            end_date="" if ongoing else _tidy_date(match.group("end")),
            current=ongoing,
        )

        parts = _chunks(_strip_dates(text)) or pending
        if parts and _looks_like_location(parts[-1]):
            record.location = parts[-1]
            parts = parts[:-1]

        if parts:
            # "Artificial Intelligence Engineer, HCLTech" -- the title comes
            # first and everything after the first comma is the employer, which
            # can itself have commas in it: "DRDO - Ministry of Defence, India".
            head, _, tail = parts[0].partition(", ")
            if head:
                record.title, record.company = head.strip(), tail.strip()
            else:
                record.title = parts[0]
            extra = [p for p in parts[1:] if not _looks_like_location(p)]
            if extra and not record.company:
                record.company = extra[0]
            elif extra:
                record.company = f"{record.company}, {extra[0]}"

        records.append(record)
        pending = []

    return records


def _parse_education(lines: list[Line]) -> list[EducationRecord]:
    """One record per line naming a degree or an institution.

    A line that merely mentions a school word in passing -- a thesis title about
    "College-Level" questions, say -- is not an entry.
    """
    records: list[EducationRecord] = []

    for line in lines:
        text = line.text.strip()
        if line.bullet:
            continue

        parts = _chunks(_strip_dates(text))
        location = ""
        if parts and _looks_like_location(parts[-1]):
            location = parts[-1]
            parts = parts[:-1]

        pieces = [p.strip() for chunk in parts for p in chunk.split(", ") if p.strip()]
        degree_piece = next((p for p in pieces if _DEGREE.search(p)), "")
        school_piece = next((p for p in pieces if _SCHOOL.search(p)), "")

        if not degree_piece and not school_piece:
            if records and re.search(r"\b(?:gpa|cgpa)\b", text, re.IGNORECASE):
                gpa = re.search(r"\b(?:gpa|cgpa)\b[:\s]*([0-9](?:\.[0-9]{1,2})?)", text, re.I)
                if gpa:
                    records[-1].gpa = gpa.group(1)
            continue

        if not school_piece and degree_piece:
            # "B.Tech. in Mechanical Engineering, IIITDM Kurnool" -- the school
            # is not always spelled with a word like "University" in it.
            remaining = [p for p in pieces if p != degree_piece]
            school_piece = remaining[0] if remaining else ""

        record = EducationRecord(school=school_piece, location=location)

        if degree_piece:
            degree_match = _DEGREE.search(degree_piece)
            record.degree = degree_match.group(0).strip() if degree_match else degree_piece
            study = _FIELD_OF_STUDY.search(degree_piece)
            if study:
                record.field_of_study = study.group(1).strip(" ,.")

        date_match = _RANGE.search(text)
        if date_match:
            record.start_date = _tidy_date(date_match.group("start"))
            record.end_date = _tidy_date(date_match.group("end"))
        else:
            year = re.search(r"(?:19|20)\d{2}", text)
            if year:
                record.end_date = year.group(0)

        records.append(record)

    return records


def _parse_skills(lines: list[Line]) -> list[str]:
    skills: list[str] = []
    for line in lines:
        body = re.sub(r"^[^:]{0,40}:", "", line.text)
        for part in re.split(r"[,;|·•]", body):
            skill = part.strip(" .-–—\t")
            if 1 < len(skill) <= 40 and not skill.endswith("."):
                skills.append(skill)
    seen: set[str] = set()
    unique = []
    for skill in skills:
        key = skill.lower()
        if key not in seen:
            seen.add(key)
            unique.append(skill)
    return unique[:60]


def _guess_name(header: list[str]) -> str:
    for line in header[:4]:
        if _EMAIL.search(line) or _PHONE.search(line) or "http" in line.lower():
            continue
        words = line.split()
        if 1 < len(words) <= 5 and all(re.fullmatch(r"[A-Za-z.'-]+", w) for w in words):
            return line.strip().title() if line.isupper() else line.strip()
    return ""


def _guess_location(header: list[str]) -> str:
    for line in header[:5]:
        for chunk in _chunks(line) + [p.strip() for p in line.split("|")]:
            if _looks_like_location(chunk):
                return chunk
    return ""


def extract(source: Path | bytes) -> ResumeExtract:
    """Read a resume into records. Anything not stated comes back empty."""
    parsed = read_docx(source)
    lines = [line.text.strip() for line in parsed]
    sections = _split_sections(parsed)
    header = [line.text.strip() for line in sections.get("header", [])]
    blob = "\n".join(lines)

    result = ResumeExtract(lines=lines)
    result.name = _guess_name(header)
    result.location = _guess_location(header)

    email = _EMAIL.search(blob)
    result.email = email.group(0) if email else ""
    linkedin = _LINKEDIN.search(blob)
    result.linkedin = _with_scheme(linkedin.group(0)) if linkedin else ""
    github = _GITHUB.search(blob)
    result.github = _with_scheme(github.group(0)) if github else ""

    header_blob = "\n".join(header)
    phone = _PHONE.search(header_blob) or _PHONE.search(blob)
    result.phone = re.sub(r"\s+", " ", phone.group(0)).strip() if phone else ""

    for candidate in _WEBSITE.finditer(header_blob):
        url = candidate.group(0)
        if "linkedin.com" in url or "github.com" in url:
            continue
        result.website = url
        break

    result.experience = _parse_experience(sections.get("experience", []))
    result.education = _parse_education(sections.get("education", []))
    result.skills = _parse_skills(sections.get("skills", []))

    if not result.experience:
        result.notes.append(
            "no dated employment entries were found; add them by hand in the panel"
        )
    if not result.education:
        result.notes.append("no education entries were found; add them by hand in the panel")
    return result


def _with_scheme(url: str) -> str:
    return url if url.lower().startswith("http") else "https://" + url.lstrip("/")
