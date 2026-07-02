from io import BytesIO

from docx import Document
from pypdf import PdfReader

from applypilot.documents import (
    artifact_filename,
    build_cover_letter_docx,
    build_docx,
    build_pdf,
    build_reconstructed_resume_docx,
)
from applypilot.models import (
    CandidateProfile,
    GeneratedCoverLetter,
    JobContext,
    ResumeDocument,
    TailoredArtifact,
    TailoredBullet,
    TailoredExperience,
    TailoredResume,
)


def sample_artifact() -> TailoredArtifact:
    return TailoredArtifact(
        tailored=TailoredResume(
            headline="Software Engineer | Python Automation",
            summary="Software engineer focused on reliable automation and APIs.",
            skills=["Python", "SQL", "FastAPI", "Automated Testing"],
            experiences=[
                TailoredExperience(
                    heading="Software Engineer - Example Systems",
                    bullets=[
                        TailoredBullet(
                            text="Built reliable Python services and deployment automation.",
                            evidence_ids=["evidence-1"],
                        ),
                        TailoredBullet(
                            text="Created automated tests for critical API workflows.",
                            evidence_ids=["evidence-2"],
                        ),
                    ],
                )
            ],
        )
    )


def sample_profile() -> CandidateProfile:
    return CandidateProfile(
        legal_name="Test Candidate",
        email="candidate@example.test",
        phone="+1 555 0100",
        city="Chicago",
        region="IL",
        country="US",
        linkedin_url="https://www.linkedin.com/in/test-candidate",
    )


def test_build_docx_contains_tailored_content() -> None:
    content = build_docx(sample_artifact(), sample_profile())
    document = Document(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert content.startswith(b"PK")
    assert "Test Candidate" in text
    assert "PROFESSIONAL SUMMARY" in text
    assert "Built reliable Python services" in text


def test_build_pdf_is_single_page_and_extractable() -> None:
    content = build_pdf(sample_artifact(), sample_profile())
    reader = PdfReader(BytesIO(content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)

    assert content.startswith(b"%PDF")
    assert len(reader.pages) == 1
    assert "Test Candidate" in text
    assert "Python Automation" in text


def test_artifact_filename_is_safe() -> None:
    assert artifact_filename(sample_profile(), "pdf") == "test-candidate-tailored-resume.pdf"


def test_reconstructed_resume_and_generated_cover_letter_are_attachable_docx() -> None:
    resume = ResumeDocument(
        filename="legacy.pdf",
        media_type="application/pdf",
        sha256="legacy",
        extracted_text="TEST CANDIDATE\nEXPERIENCE\nBuilt reliable Python services.",
    )
    reconstructed = Document(BytesIO(build_reconstructed_resume_docx(resume)))
    letter = GeneratedCoverLetter(
        job_title="Software Engineer",
        company="Example Robotics",
        body="Dear Hiring Team,\n\nI offer verified Python automation experience.\n\nSincerely,",
    )
    cover = Document(
        BytesIO(
            build_cover_letter_docx(
                letter,
                sample_profile(),
                JobContext(
                    title="Software Engineer",
                    company="Example Robotics",
                    description="Build automation.",
                ),
            )
        )
    )

    assert "Built reliable Python services" in "\n".join(
        item.text for item in reconstructed.paragraphs
    )
    assert "verified Python automation" in "\n".join(item.text for item in cover.paragraphs)
