from io import BytesIO

import pytest
from docx import Document

from applypilot.resume import ResumeExtractionError, extract_resume


def test_extract_txt_resume() -> None:
    text = "Software Engineer\nBuilt reliable Python services and automated deployments. " * 3

    resume = extract_resume("candidate.txt", text.encode(), "text/plain")

    assert resume.filename == "candidate.txt"
    assert "Python services" in resume.extracted_text
    assert len(resume.sha256) == 64


def test_extract_docx_includes_paragraphs_and_tables() -> None:
    document = Document()
    document.add_heading("Test Candidate", level=1)
    document.add_paragraph(
        "Software engineer who built reliable Python services and automated deployments."
    )
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, SQL, FastAPI"
    stream = BytesIO()
    document.save(stream)

    resume = extract_resume("candidate.docx", stream.getvalue())

    assert "Test Candidate" in resume.extracted_text
    assert "Skills | Python, SQL, FastAPI" in resume.extracted_text


def test_reject_unsupported_resume_type() -> None:
    with pytest.raises(ResumeExtractionError, match="DOCX, PDF, or TXT"):
        extract_resume("candidate.rtf", b"not supported")
