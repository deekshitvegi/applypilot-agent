"""The local service, end to end, against a throwaway data directory.

Nothing personal appears in this file. The resume it uploads is built here.
"""

from __future__ import annotations

import importlib
from io import BytesIO

import pytest
from fastapi.testclient import TestClient


@pytest.fixture
def client(tmp_path, monkeypatch):
    monkeypatch.setenv("APPLYPILOT_DATA_DIR", str(tmp_path / "data"))
    import applypilot.config as config
    import applypilot.main as main

    importlib.reload(config)
    importlib.reload(main)
    with TestClient(main.app) as test_client:
        yield test_client


def make_resume() -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph("Alex Rivera")
    document.add_paragraph("Austin, TX  |  alex@example.test  |  512 555 0147")
    document.add_paragraph("Professional Experience")
    document.add_paragraph("Machine Learning Engineer, Northwind Labs\tAustin, TX  |  Jun 2022 - Present")
    document.add_paragraph("Built retrieval pipelines over internal documents.", style="List Bullet")
    document.add_paragraph("Data Analyst, Contoso\tDallas, TX  |  Jan 2020 - May 2022")
    document.add_paragraph("Reported on customer churn.", style="List Bullet")
    document.add_paragraph("Education")
    document.add_paragraph("M.S. in Computer Science, University of Example\tAustin, TX  |  Aug 2018 - May 2020")
    document.add_paragraph("Technical Skills")
    document.add_paragraph("Languages: Python, SQL, TypeScript")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_health_reports_the_version_the_service_is_running(client):
    from applypilot import __version__

    body = client.get("/health").json()
    assert body["ok"] is True
    assert body["version"] == __version__
    assert body["onboarding_complete"] is False
    assert body["missing_for_applications"]


def test_the_extension_manifest_version_matches_the_service(client):
    import json
    from pathlib import Path

    from applypilot import __version__

    manifest = json.loads(
        (Path(__file__).resolve().parents[1] / "extension" / "manifest.json").read_text("utf-8")
    )
    assert manifest["version"] == __version__, (
        "the manifest and the service must ship the same version; the panel warns when "
        "they differ and this keeps that warning meaningful"
    )


def test_onboarding_starts_with_the_legal_questions_unanswered(client):
    body = client.get("/onboarding").json()
    assert body["complete"] is False
    keys = [step["key"] for step in body["steps"]]
    for key in ("work_authorization", "requires_sponsorship", "over_18", "background_check_consent"):
        assert key in keys
    assert any("come up on nearly every form" in note for note in body["notes"])


def test_answering_onboarding_sticks(client):
    client.post("/onboarding/answer", json={"key": "requires_sponsorship", "value": "No"})
    profile = client.get("/profile").json()
    assert profile["facts"]["requires_sponsorship"] == "No"


def test_uploading_a_resume_fills_structured_records(client):
    response = client.post(
        "/resume",
        files={"file": ("resume.docx", make_resume(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 200, response.text
    body = response.json()

    assert [e["school"] for e in body["education"]] == ["University of Example"]
    assert body["education"][0]["degree"] == "M.S."
    assert body["education"][0]["field_of_study"] == "Computer Science"

    companies = [x["company"] for x in body["experience"]]
    assert companies == ["Northwind Labs", "Contoso"]
    assert body["experience"][0]["current"] is True

    profile = client.get("/profile").json()
    assert profile["facts"]["full_name"] == "Alex Rivera"
    assert profile["facts"]["state"] == "Texas"


def test_a_non_docx_upload_is_refused_with_a_reason(client):
    response = client.post("/resume", files={"file": ("resume.pdf", b"%PDF-1.4", "application/pdf")})
    assert response.status_code == 400
    assert ".docx" in response.json()["detail"]


def test_planning_an_application_fills_what_it_can_and_asks_the_rest(client):
    client.post(
        "/resume",
        files={"file": ("resume.docx", make_resume(), "application/octet-stream")},
    )
    for key, value in [
        ("requires_sponsorship", "No"),
        ("work_authorization", "Yes"),
        ("country", "United States"),
    ]:
        client.post("/onboarding/answer", json={"key": key, "value": value})

    observation = {
        "url": "https://boards.greenhouse.io/acme/jobs/1",
        "kind": "application",
        "signature": "sig-1",
        "fields": [
            {"fingerprint": "f1", "label": "First Name", "control": "text", "required": True},
            {"fingerprint": "f2", "label": "Email", "control": "email", "required": True},
            {
                "fingerprint": "f3",
                "label": "Will you now or in the future require visa sponsorship?",
                "control": "select",
                "required": True,
                "options": [{"label": "Yes"}, {"label": "No"}],
                "options_source": "native",
            },
            {
                "fingerprint": "f4",
                "label": "Describe a project you are proud of",
                "control": "textarea",
                "required": True,
            },
            {"fingerprint": "f5", "label": "Middle Name", "control": "text"},
        ],
    }
    body = client.post("/plan", json=observation).json()

    assert body["kind"] == "application"
    assert body["adapter"] == "greenhouse"
    assert body["host_role"] == "employer"

    values = {a["fingerprint"]: a["value"] for a in body["actions"]}
    assert values["f2"] == "alex@example.test"
    assert values["f3"] == "No"

    asked = {q["label"] for q in body["questions"]}
    assert "Describe a project you are proud of" in asked
    assert "Middle Name" not in asked, "an optional extra is left blank, not asked about"

    states = {item["label"]: item["state"] for item in body["checklist"]}
    assert states["Middle Name"] == "skipped"


def test_a_search_page_produces_no_actions_at_all(client):
    body = client.post(
        "/plan",
        json={"url": "https://www.dice.com/jobs?q=ml", "kind": "search", "fields": []},
    ).json()
    assert body["actions"] == []
    assert body["questions"] == []
    assert "nothing here to fill in" in body["narration"]


def test_results_keep_a_failure_and_report_it(client):
    payload = {
        "observation": {
            "url": "https://boards.greenhouse.io/acme/jobs/1",
            "kind": "application",
            "fields": [{"fingerprint": "f1", "label": "Country", "control": "combobox"}],
        },
        "results": [
            {
                "fingerprint": "f1",
                "label": "Country",
                "requested": "United States",
                "outcome": "failed",
                "signal": "hidden_backing_input",
                "evidence": "the page holds nothing for this control",
            },
            {
                "fingerprint": "f1",
                "label": "Country",
                "requested": "United States",
                "outcome": "verified",
                "signal": "none",
                "evidence": "read back the text the executor typed",
            },
        ],
    }
    body = client.post("/results", json=payload).json()
    assert body["failed"], "a failure must not be overwritten by an unverifiable success"
    assert body["failed"][0]["label"] == "Country"
    assert "1 failed" in body["summary"]


def test_options_from_a_control_with_no_popup_are_refused(client):
    body = client.post(
        "/options",
        json={"fingerprint": "f1", "saved_value": "United States", "options": [], "source": "none"},
    ).json()
    assert body["chosen"] is None
    assert "no list of its own" in body["note"]


def test_options_from_an_owned_popup_are_ranked(client):
    body = client.post(
        "/options",
        json={
            "fingerprint": "f1",
            "fact_key": "country",
            "saved_value": "United States",
            "source": "owned_popup",
            "options": [
                {"label": "United States Minor Outlying Islands"},
                {"label": "United States"},
                {"label": "Canada"},
            ],
        },
    ).json()
    assert body["chosen"] == "United States"


def test_learning_refuses_an_option_id(client):
    body = client.post(
        "/learn",
        json={
            "field": {
                "fingerprint": "f1",
                "label": "Country",
                "control": "select",
                "options": [{"label": "United States"}],
                "options_source": "native",
            },
            "value": "28468",
        },
    ).json()
    assert body["learned"] is False
    assert "option id" in body["reason"]
    assert client.get("/learned").json()["answers"] == []


def test_learning_keeps_a_real_answer_and_can_forget_it(client):
    field = {
        "fingerprint": "f1",
        "label": "What is your preferred start date?",
        "control": "text",
    }
    assert client.post("/learn", json={"field": field, "value": "Two weeks"}).json()["learned"]
    assert client.get("/learned").json()["answers"][0]["value"] == "Two weeks"
    client.request("DELETE", "/learned", params={"question": field["label"]})
    assert client.get("/learned").json()["answers"] == []


def test_chat_changes_a_field_and_says_so(client):
    body = client.post(
        "/chat",
        json={
            "text": "state texas",
            "fields": [
                {
                    "fingerprint": "f1",
                    "label": "State",
                    "control": "select",
                    "options": [{"label": "Texas"}, {"label": "Utah"}],
                    "options_source": "native",
                }
            ],
        },
    ).json()
    assert body["kind"] == "action"
    assert body["action"]["option_label"] == "Texas"
    assert client.get("/profile").json()["facts"]["state"] == "Texas"


def test_routing_never_stops_on_a_job_board(client):
    body = client.post(
        "/route",
        json={
            "url": "https://www.linkedin.com/jobs/view/1",
            "kind": "listing",
            "company": "Acme",
            "title": "ML Engineer",
        },
    ).json()
    assert body["action"] != "stop"
    assert body["host_role"] == "board"


def test_sign_in_details_are_never_released_for_a_registration_page(client):
    client.post("/sign-in/authorise", json={"host": "workforcenow.adp.com"})
    body = client.post(
        "/sign-in/check",
        json={"url": "https://workforcenow.adp.com/register", "kind": "registration"},
    ).json()
    assert body["allowed"] is False
    assert "accepts the employer's terms" in body["reason"]


def test_sign_in_details_are_never_released_to_a_lookalike_host(client):
    client.post("/sign-in/authorise", json={"host": "workforcenow.adp.com"})
    body = client.post(
        "/sign-in/check",
        json={"url": "https://workforcenow.adp.com.evil.test/login", "kind": "sign_in"},
    ).json()
    assert body["allowed"] is False


def test_sign_in_details_are_released_only_once_for_the_exact_host(client):
    client.post("/sign-in/authorise", json={"host": "workforcenow.adp.com"})
    first = client.post(
        "/sign-in/check",
        json={"url": "https://workforcenow.adp.com/mascsr/login", "kind": "sign_in"},
    ).json()
    second = client.post(
        "/sign-in/check",
        json={"url": "https://workforcenow.adp.com/mascsr/login", "kind": "sign_in"},
    ).json()
    assert first["allowed"] is True
    assert second["allowed"] is False


def test_an_application_is_submitted_only_with_a_confirmation(client):
    created = client.post(
        "/applications",
        json={"company": "Acme", "role": "ML Engineer", "url": "https://boards.greenhouse.io/acme/jobs/1"},
    ).json()

    without = client.post(
        "/applications/submitted", json={"application_id": created["id"], "confirmation": ""}
    ).json()
    assert without["application"]["status"] == "ready_to_submit"
    assert "not recorded this as submitted" in without["message"]

    with_confirmation = client.post(
        "/applications/submitted",
        json={"application_id": created["id"], "confirmation": "Thank you for applying"},
    ).json()
    assert with_confirmation["application"]["status"] == "submitted"
    assert with_confirmation["application"]["applied_on"]


def test_applications_export_as_csv(client):
    client.post(
        "/applications",
        json={"company": "Acme", "role": "ML Engineer", "url": "https://example.test/1"},
    )
    text = client.get("/applications/export").text
    assert text.splitlines()[0].startswith("applied_on,company,role,status")
    assert "Acme" in text


def test_a_tailored_resume_only_reorders_what_is_already_there(client):
    client.post("/resume", files={"file": ("resume.docx", make_resume(), "application/octet-stream")})
    body = client.post(
        "/documents/tailored-resume",
        json={
            "job_title": "Machine Learning Engineer",
            "company": "Acme",
            "job_description": "retrieval pipelines, Python, embeddings",
        },
    ).json()
    assert body["document"]["kind"] == "tailored_resume"
    assert body["ordering"][0].startswith("Machine Learning Engineer at Northwind Labs")
    assert any("nothing was added" in note for note in body["notes"])


def test_the_api_key_is_stored_but_never_read_back(client):
    client.put("/settings", json={"model_api_key": "test-key-not-a-real-one"})
    body = client.get("/settings").json()
    assert body["model_configured"] is True
    assert "test-key-not-a-real-one" not in str(body)


def test_a_stored_document_can_be_fetched_for_attaching(client):
    # Built once: a .docx is a zip, and rebuilding it gives different bytes.
    document = make_resume()
    upload = client.post(
        "/resume", files={"file": ("resume.docx", document, "application/octet-stream")}
    ).json()
    document_id = upload["document"]["id"]

    body = client.get(f"/documents/{document_id}/content").json()
    assert body["filename"] == "resume.docx"
    assert body["mime"].endswith("wordprocessingml.document")

    import base64

    assert base64.b64decode(body["base64"]) == document, "the bytes come back intact"
    assert client.get("/documents/nope/content").status_code == 404


def test_attaching_the_resume_is_on_by_default_and_can_be_turned_off(client):
    assert client.get("/settings").json()["auto_attach_resume"] is True
    client.put("/settings", json={"auto_attach_resume": False})
    assert client.get("/settings").json()["auto_attach_resume"] is False


def test_history_answers_land_in_the_record_they_belong_to(client):
    client.post("/resume", files={"file": ("resume.docx", make_resume(), "application/octet-stream")})
    client.post("/profile/fact", json={"fact_key": "education.gpa", "value": "3.8", "entry": 0})

    profile = client.get("/profile").json()
    assert profile["education"][0]["gpa"] == "3.8"
    assert "education.gpa" not in profile["facts"], "not into the flat facts, where nothing reads it"


# ---------------------------------------------------------------------------
# 71. Only an attempt to move on can stall a page.
# ---------------------------------------------------------------------------


def _blank_application() -> dict:
    return {
        "url": "https://careers.example.test/apply",
        "kind": "application",
        "signature": "one-and-the-same",
        "fields": [],
    }


def test_planning_the_same_page_over_and_over_is_not_a_stall(client):
    """Filling a page re-plans it several times over, by design.

    The page is read again after every choice and again on each correction
    pass. Counting those as failures to progress declared every page stuck on
    the third look -- which is how a run reached a new step, said it had
    stopped, and then sat there with a plan it would not act on.
    """
    stuck = "This page has not changed after several attempts"
    for _ in range(6):
        notes = client.post("/plan", json=_blank_application()).json()["notes"]
        assert not any(stuck in note for note in notes), notes


def test_pressing_continue_and_getting_nowhere_still_stalls(client):
    """The guard itself is untouched: it just counts the right thing now."""
    stuck = "This page has not changed after several attempts"
    seen = []
    for _ in range(5):
        response = client.post(
            "/plan", params={"after_continue": "true"}, json=_blank_application()
        )
        seen.append(any(stuck in note for note in response.json()["notes"]))
    assert seen[-1] is True, seen
